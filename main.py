import sys
import fitz  # PyMuPDF
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QPushButton, QLabel, QTableWidget, 
                             QSplitter, QTextEdit, QTabWidget, QLineEdit, 
                             QFormLayout, QHeaderView, QFileDialog,
                             QGraphicsView, QGraphicsScene, QGraphicsPixmapItem,
                             QDialog, QTableWidgetItem, QMessageBox,
                             QDateEdit, QRadioButton, QButtonGroup, QPlainTextEdit)
from PyQt6.QtCore import Qt, pyqtSignal, QRectF, QTimer, QDate, QMarginsF
from PyQt6.QtGui import QPdfWriter, QTextDocument, QPageSize, QPageLayout, QPixmap, QImage

import database as db
import validador_citacoes
import os
import smtplib
from email.message import EmailMessage
from dotenv import load_dotenv

load_dotenv()
DARK_GLASS_QSS = """
QWidget {
    background-color: #0b0f19;
    color: #e2e8f0;
    font-family: "Segoe UI", "Inter", sans-serif;
    font-size: 13px;
}

QDialog, QMainWindow {
    background-color: #0b0f19;
}

QLabel {
    background: transparent;
}

/* Simulated Frosted Glass Panels */
QTableWidget, QScrollArea, QTabWidget::pane {
    background-color: rgba(255, 255, 255, 0.03);
    border: 1px solid rgba(255, 255, 255, 0.08);
    border-radius: 8px;
}

/* Table Specifics */
QTableWidget {
    gridline-color: rgba(255, 255, 255, 0.05);
    selection-background-color: rgba(59, 130, 246, 0.3);
    selection-color: #ffffff;
}
QHeaderView::section {
    background-color: rgba(0, 0, 0, 0.3);
    color: #94a3b8;
    border: none;
    padding: 6px;
    font-weight: bold;
}
QTableCornerButton::section {
    background-color: rgba(0, 0, 0, 0.3);
}

/* Inputs */
QLineEdit, QTextEdit, QPlainTextEdit {
    background-color: rgba(0, 0, 0, 0.3);
    border: 1px solid rgba(255, 255, 255, 0.1);
    border-radius: 6px;
    padding: 8px;
    color: #e2e8f0;
}
QLineEdit:focus, QTextEdit:focus, QPlainTextEdit:focus {
    border: 1px solid rgba(59, 130, 246, 0.6);
    background-color: rgba(0, 0, 0, 0.4);
}

/* Buttons */
QPushButton {
    background-color: rgba(255, 255, 255, 0.08);
    border: 1px solid rgba(255, 255, 255, 0.15);
    border-radius: 6px;
    color: #ffffff;
    padding: 8px 16px;
    font-weight: bold;
}
QPushButton:hover {
    background-color: rgba(255, 255, 255, 0.15);
    border: 1px solid rgba(255, 255, 255, 0.3);
}
QPushButton:pressed {
    background-color: rgba(255, 255, 255, 0.05);
}

/* Colored Buttons */
QPushButton#btn_green {
    background-color: rgba(16, 185, 129, 0.8);
    border: 1px solid #10b981;
}
QPushButton#btn_green:hover { background-color: rgba(16, 185, 129, 1.0); }

QPushButton#btn_blue {
    background-color: rgba(59, 130, 246, 0.8);
    border: 1px solid #3b82f6;
}
QPushButton#btn_blue:hover { background-color: rgba(59, 130, 246, 1.0); }

QPushButton#btn_red {
    background-color: rgba(239, 68, 68, 0.8);
    border: 1px solid #ef4444;
}
QPushButton#btn_red:hover { background-color: rgba(239, 68, 68, 1.0); }

QPushButton#btn_orange {
    background-color: rgba(249, 115, 22, 0.8);
    border: 1px solid #f97316;
}
QPushButton#btn_orange:hover { background-color: rgba(249, 115, 22, 1.0); }

/* Icon Buttons */
QPushButton#btn_icon_blue, QPushButton#btn_icon_orange {
    padding: 0px;
    font-size: 16px;
    min-width: 36px;
    max-width: 36px;
    min-height: 36px;
    max-height: 36px;
    border-radius: 6px;
}
QPushButton#btn_icon_blue {
    background-color: rgba(59, 130, 246, 0.8);
    border: 1px solid #3b82f6;
}
QPushButton#btn_icon_blue:hover { background-color: rgba(59, 130, 246, 1.0); }
QPushButton#btn_icon_orange {
    background-color: rgba(249, 115, 22, 0.8);
    border: 1px solid #f97316;
}
QPushButton#btn_icon_orange:hover { background-color: rgba(249, 115, 22, 1.0); }

/* Tabs */
QTabBar::tab {
    background: rgba(255, 255, 255, 0.05);
    border: 1px solid rgba(255, 255, 255, 0.1);
    border-bottom: none;
    border-top-left-radius: 6px;
    border-top-right-radius: 6px;
    padding: 8px 16px;
    margin-right: 2px;
    color: #94a3b8;
}
QTabBar::tab:selected {
    color: #ffffff;
}

/* Splitter */
QSplitter::handle {
    background: rgba(255, 255, 255, 0.1);
    width: 2px;
}
"""

class NewProjectDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Novo Projeto de Revisão")
        self.setMinimumWidth(400)
        
        self.student_name = ""
        self.thesis_title = ""
        self.pdf_path = ""
        
        layout = QVBoxLayout(self)
        form = QFormLayout()
        
        self.input_name = QLineEdit()
        self.input_title = QLineEdit()
        self.input_advisor_name = QLineEdit()
        self.input_advisor_email = QLineEdit()
        
        form.addRow("Nome do Aluno:", self.input_name)
        form.addRow("Título do TCC:", self.input_title)
        form.addRow("Nome do Orientador:", self.input_advisor_name)
        form.addRow("E-mail do Orientador:", self.input_advisor_email)
        
        self.btn_select_pdf = QPushButton("Selecionar Arquivo PDF")
        self.btn_select_pdf.setObjectName("btn_blue")
        self.btn_select_pdf.clicked.connect(self.select_pdf)
        self.lbl_pdf_path = QLabel("Nenhum PDF selecionado")
        self.lbl_pdf_path.setStyleSheet("color: gray;")
        
        form.addRow("Arquivo PDF:", self.btn_select_pdf)
        form.addRow("", self.lbl_pdf_path)
        
        layout.addLayout(form)
        
        btn_layout = QHBoxLayout()
        btn_create = QPushButton("Criar Projeto")
        btn_create.setObjectName("btn_green")
        btn_cancel = QPushButton("Cancelar")
        btn_cancel.setObjectName("btn_red")
        
        btn_create.clicked.connect(self.accept_creation)
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_create)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

    def select_pdf(self):
        path, _ = QFileDialog.getOpenFileName(self, "Abrir PDF do TCC", "", "Arquivos PDF (*.pdf)")
        if path:
            self.pdf_path = path
            self.lbl_pdf_path.setText(path)

    def accept_creation(self):
        self.student_name = self.input_name.text().strip()
        self.thesis_title = self.input_title.text().strip()
        self.advisor_name = self.input_advisor_name.text().strip()
        self.advisor_email = self.input_advisor_email.text().strip()
        
        if not self.student_name or not self.thesis_title or not self.pdf_path:
            QMessageBox.warning(self, "Aviso", "Por favor, preencha todos os campos obrigatórios (Aluno, TCC) e selecione um PDF.")
            return
            
        self.accept()

class EditProjectDialog(QDialog):
    def __init__(self, project_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Editar Projeto de Revisão")
        self.setMinimumWidth(400)
        
        self.project_id = project_data['id']
        self.student_name = project_data['student_name']
        self.thesis_title = project_data['thesis_title']
        self.advisor_name = project_data.get('advisor_name', '')
        self.advisor_email = project_data.get('advisor_email', '')
        
        layout = QVBoxLayout(self)
        form = QFormLayout()
        
        self.input_name = QLineEdit(self.student_name)
        self.input_title = QLineEdit(self.thesis_title)
        self.input_advisor_name = QLineEdit(self.advisor_name)
        self.input_advisor_email = QLineEdit(self.advisor_email)
        
        form.addRow("Nome do Aluno:", self.input_name)
        form.addRow("Título do TCC:", self.input_title)
        form.addRow("Nome do Orientador:", self.input_advisor_name)
        form.addRow("E-mail do Orientador:", self.input_advisor_email)
        
        layout.addLayout(form)
        
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Salvar Alterações")
        btn_save.setObjectName("btn_green")
        btn_cancel = QPushButton("Cancelar")
        btn_cancel.setObjectName("btn_red")
        
        btn_save.clicked.connect(self.accept_edit)
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

    def accept_edit(self):
        self.student_name = self.input_name.text().strip()
        self.thesis_title = self.input_title.text().strip()
        self.advisor_name = self.input_advisor_name.text().strip()
        self.advisor_email = self.input_advisor_email.text().strip()
        
        if not self.student_name or not self.thesis_title:
            QMessageBox.warning(self, "Aviso", "Por favor, preencha nome do aluno e título do TCC.")
            return
            
        self.accept()

class ExportDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Opções de Exportação")
        self.setMinimumWidth(400)
        
        self.evaluator_name = ""
        self.date_str = ""
        self.export_format = "pdf"
        
        layout = QVBoxLayout(self)
        form = QFormLayout()
        
        self.input_evaluator = QLineEdit()
        self.input_evaluator.setPlaceholderText("Ex: Prof. Dr. João Silva")
        
        self.input_date = QDateEdit()
        self.input_date.setCalendarPopup(True)
        self.input_date.setDate(QDate.currentDate())
        
        form.addRow("Avaliador(a):", self.input_evaluator)
        form.addRow("Data:", self.input_date)
        
        layout.addLayout(form)
        
        format_layout = QHBoxLayout()
        self.radio_pdf = QRadioButton("PDF (.pdf)")
        self.radio_pdf.setChecked(True)
        self.radio_docx = QRadioButton("Word (.docx)")
        
        self.format_group = QButtonGroup()
        self.format_group.addButton(self.radio_pdf)
        self.format_group.addButton(self.radio_docx)
        
        format_layout.addWidget(self.radio_pdf)
        format_layout.addWidget(self.radio_docx)
        
        form.addRow("Formato:", format_layout)
        
        btn_layout = QHBoxLayout()
        btn_export = QPushButton("Exportar")
        btn_export.setObjectName("btn_green")
        btn_cancel = QPushButton("Cancelar")
        btn_cancel.setObjectName("btn_red")
        
        btn_export.clicked.connect(self.accept_export)
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_export)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

    def accept_export(self):
        self.evaluator_name = self.input_evaluator.text().strip()
        self.date_str = self.input_date.date().toString("dd/MM/yyyy")
        self.export_format = "pdf" if self.radio_pdf.isChecked() else "docx"
        self.accept()

class PDFViewer(QGraphicsView):
    textSelected = pyqtSignal(str, int)  # Emits captured text and page number

    def __init__(self):
        super().__init__()
        self.scene = QGraphicsScene(self)
        self.setScene(self.scene)
        # Enable Rubber Band dragging for selecting area
        self.setDragMode(QGraphicsView.DragMode.RubberBandDrag)
        
        self.pdf_document = None
        self.current_page_idx = 0
        self.pixmap_item = None
        self.zoom_factor = 1.25
        self.current_scale = 1.0

    def load_pdf(self, file_path):
        try:
            self.pdf_document = fitz.open(file_path)
            self.current_page_idx = 0
            self.current_scale = 1.0
            self.render_page()
            return True
        except Exception as e:
            print(f"Erro ao carregar o PDF: {e}")
            return False

    def render_page(self):
        if not self.pdf_document: return
        
        page = self.pdf_document[self.current_page_idx]
        
        # Create a transform matrix to scale the rendering resolution
        mat = fitz.Matrix(self.current_scale, self.current_scale)
        pix = page.get_pixmap(matrix=mat)
        
        # Convert PyMuPDF Pixmap to PyQt6 QImage then QPixmap
        fmt = QImage.Format.Format_RGB888 if pix.alpha == 0 else QImage.Format.Format_RGBA8888
        img = QImage(pix.samples, pix.width, pix.height, pix.stride, fmt)
        qpixmap = QPixmap.fromImage(img)
        
        self.scene.clear()
        self.pixmap_item = QGraphicsPixmapItem(qpixmap)
        self.scene.addItem(self.pixmap_item)
        self.setSceneRect(QRectF(self.pixmap_item.pixmap().rect()))
        
    def next_page(self):
        if self.pdf_document and self.current_page_idx < len(self.pdf_document) - 1:
            self.current_page_idx += 1
            self.render_page()

    def prev_page(self):
        if self.pdf_document and self.current_page_idx > 0:
            self.current_page_idx -= 1
            self.render_page()

    def zoom_in(self):
        self.current_scale *= self.zoom_factor
        self.render_page()

    def zoom_out(self):
        self.current_scale /= self.zoom_factor
        self.render_page()

    def mouseReleaseEvent(self, event):
        # Pegar o retângulo ANTES de chamar super() pois o evento pode limpar a seleção
        rubber_band_rect = self.rubberBandRect()
        
        super().mouseReleaseEvent(event)
        
        if not self.pdf_document: return
        
        if rubber_band_rect.isEmpty():
            return
            
        scene_rect = self.mapToScene(rubber_band_rect).boundingRect()
        
        pdf_rect = fitz.Rect(
            scene_rect.left() / self.current_scale,
            scene_rect.top() / self.current_scale,
            scene_rect.right() / self.current_scale,
            scene_rect.bottom() / self.current_scale
        )
        
        page = self.pdf_document[self.current_page_idx]
        
        text = page.get_textbox(pdf_rect).strip().replace('\n', ' ')
        
        if text:
            self.textSelected.emit(text, self.current_page_idx + 1)


class EditCitationDialog(QDialog):
    def __init__(self, row_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Editar Verificação")
        self.setMinimumWidth(500)
        
        self.row_data = row_data
        
        layout = QVBoxLayout(self)
        form = QFormLayout()
        
        self.input_status = QLineEdit(str(row_data.get('Status', '')))
        self.input_citacao = QLineEdit(str(row_data.get('Citação Encontrada', '')))
        self.input_pagina = QLineEdit(str(row_data.get('Página', '')))
        
        self.input_contexto = QPlainTextEdit(str(row_data.get('Contexto', '')))
        self.input_contexto.setMaximumHeight(80)
        
        self.input_referencia = QPlainTextEdit(str(row_data.get('Referência Correspondente', '')))
        self.input_referencia.setMaximumHeight(80)
        
        form.addRow("Status:", self.input_status)
        form.addRow("Citação:", self.input_citacao)
        form.addRow("Página:", self.input_pagina)
        form.addRow("Contexto:", self.input_contexto)
        form.addRow("Referência:", self.input_referencia)
        
        layout.addLayout(form)
        
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Salvar")
        btn_save.setObjectName("btn_green")
        btn_save.clicked.connect(self.accept)
        
        btn_cancel = QPushButton("Cancelar")
        btn_cancel.setObjectName("btn_red")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addStretch()
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)
        
    def get_updated_data(self):
        return {
            'Status': self.input_status.text(),
            'Citação Encontrada': self.input_citacao.text(),
            'Página': self.input_pagina.text(),
            'Contexto': self.input_contexto.toPlainText(),
            'Referência Correspondente': self.input_referencia.toPlainText()
        }

class CitationReportDialog(QDialog):
    def __init__(self, results, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Relatório de Validação de Citações")
        self.setMinimumSize(900, 600)
        self.results = results
        
        layout = QVBoxLayout(self)
        
        lbl_info = QLabel(f"<b>{len(results)}</b> resultados encontrados.")
        layout.addWidget(lbl_info)
        
        self.table = QTableWidget(0, 5)
        self.table.setHorizontalHeaderLabels(["Status", "Citação Encontrada", "Página", "Contexto", "Referência Correspondente"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        
        layout.addWidget(self.table)
        
        self.populate_table()
        
        btn_layout = QHBoxLayout()
        
        btn_edit = QPushButton("Editar")
        btn_edit.setObjectName("btn_blue")
        btn_edit.clicked.connect(self.edit_selected)
        
        btn_delete = QPushButton("Excluir")
        btn_delete.setObjectName("btn_red")
        btn_delete.clicked.connect(self.delete_selected)
        
        btn_export = QPushButton("Exportar para Excel")
        btn_export.setObjectName("btn_green")
        btn_export.clicked.connect(self.export_excel)
        
        btn_close = QPushButton("Fechar")
        btn_close.clicked.connect(self.accept)
        
        btn_layout.addWidget(btn_edit)
        btn_layout.addWidget(btn_delete)
        btn_layout.addStretch()
        btn_layout.addWidget(btn_export)
        btn_layout.addWidget(btn_close)
        
        layout.addLayout(btn_layout)
        
    def populate_table(self):
        self.table.setRowCount(0)
        for i, res in enumerate(self.results):
            self.table.insertRow(i)
            self.table.setItem(i, 0, QTableWidgetItem(str(res.get('Status', ''))))
            self.table.setItem(i, 1, QTableWidgetItem(str(res.get('Citação Encontrada', ''))))
            self.table.setItem(i, 2, QTableWidgetItem(str(res.get('Página', ''))))
            
            context = str(res.get('Contexto', ''))
            if len(context) > 60:
                context = context[:57] + "..."
            self.table.setItem(i, 3, QTableWidgetItem(context))
            
            ref = str(res.get('Referência Correspondente', ''))
            if len(ref) > 60:
                ref = ref[:57] + "..."
            self.table.setItem(i, 4, QTableWidgetItem(ref))

    def edit_selected(self):
        selected = self.table.selectedItems()
        if not selected:
            QMessageBox.warning(self, "Aviso", "Selecione uma linha para editar.")
            return
            
        row = selected[0].row()
        row_data = self.results[row]
        
        dialog = EditCitationDialog(row_data, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            updated_data = dialog.get_updated_data()
            self.results[row] = updated_data
            
            self.table.setItem(row, 0, QTableWidgetItem(str(updated_data['Status'])))
            self.table.setItem(row, 1, QTableWidgetItem(str(updated_data['Citação Encontrada'])))
            self.table.setItem(row, 2, QTableWidgetItem(str(updated_data['Página'])))
            
            context = str(updated_data['Contexto'])
            self.table.setItem(row, 3, QTableWidgetItem(context[:57] + "..." if len(context) > 60 else context))
            
            ref = str(updated_data['Referência Correspondente'])
            self.table.setItem(row, 4, QTableWidgetItem(ref[:57] + "..." if len(ref) > 60 else ref))

    def delete_selected(self):
        selected = self.table.selectedItems()
        if not selected:
            QMessageBox.warning(self, "Aviso", "Selecione uma linha para excluir.")
            return
            
        row = selected[0].row()
        reply = QMessageBox.question(self, "Confirmar", "Tem certeza que deseja excluir esta verificação?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            self.results.pop(row)
            self.table.removeRow(row)

    def export_excel(self):
        path, _ = QFileDialog.getSaveFileName(self, "Salvar Relatório", "relatorio_citacoes.xlsx", "Excel Files (*.xlsx);;CSV Files (*.csv)")
        if path:
            try:
                import pandas as pd
                df = pd.DataFrame(self.results)
                if path.endswith('.csv'):
                    df.to_csv(path, index=False, sep=';', encoding='utf-8-sig')
                else:
                    df.to_excel(path, index=False)
                QMessageBox.information(self, "Sucesso", "Relatório exportado com sucesso!")
            except Exception as e:
                QMessageBox.critical(self, "Erro", f"Falha ao exportar: {e}")

class TCCAssistantApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Assistente de Revisão de TCC v1.0")
        self.setGeometry(100, 100, 1200, 800)
        
        db.init_db()
        
        self.current_project_id = None
        
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QVBoxLayout(self.central_widget)
        
        self.autosave_timer = QTimer(self)
        self.autosave_timer.setInterval(2000)
        self.autosave_timer.setSingleShot(True)
        self.autosave_timer.timeout.connect(self.save_general_opinion)
        
        self.init_dashboard()
        
    def init_dashboard(self):
        self.current_project_id = None
        self.clear_layout(self.main_layout)
        
        dashboard_widget = QWidget()
        layout = QVBoxLayout(dashboard_widget)
        
        title = QLabel("<h2>Painel de Projetos</h2>")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)
        
        btn_layout = QHBoxLayout()
        btn_new = QPushButton("Nova Revisão")
        btn_new.setObjectName("btn_green")
        btn_open = QPushButton("Abrir Revisão")
        btn_open.setObjectName("btn_blue")
        btn_edit = QPushButton("Editar")
        btn_edit.setObjectName("btn_blue")
        btn_status = QPushButton("Alterar Status")
        btn_status.setObjectName("btn_orange")
        btn_delete = QPushButton("Excluir")
        btn_delete.setObjectName("btn_red")
        
        btn_new.clicked.connect(self.create_new_project)
        btn_open.clicked.connect(self.open_selected_project)
        btn_edit.clicked.connect(self.edit_selected_project)
        btn_status.clicked.connect(self.toggle_selected_project_status)
        btn_delete.clicked.connect(self.delete_selected_project)
        
        btn_layout.addWidget(btn_new)
        btn_layout.addWidget(btn_open)
        btn_layout.addWidget(btn_edit)
        btn_layout.addWidget(btn_status)
        btn_layout.addWidget(btn_delete)
        layout.addLayout(btn_layout)
        
        self.project_list = QTableWidget(0, 5)
        self.project_list.setHorizontalHeaderLabels(["ID", "Nome do Aluno", "Título", "Status", "Última Atualização"])
        self.project_list.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.project_list.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.project_list.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.project_list.hideColumn(0)
        layout.addWidget(self.project_list)
        
        self.main_layout.addWidget(dashboard_widget)
        self.populate_dashboard()

    def populate_dashboard(self):
        self.project_list.setRowCount(0)
        projects = db.get_projects()
        for row_idx, p in enumerate(projects):
            self.project_list.insertRow(row_idx)
            self.project_list.setItem(row_idx, 0, QTableWidgetItem(str(p['id'])))
            self.project_list.setItem(row_idx, 1, QTableWidgetItem(p['student_name']))
            self.project_list.setItem(row_idx, 2, QTableWidgetItem(p['thesis_title']))
            
            # Translating default DB English status if present
            status = "Em Andamento" if p['status'] == "In Progress" else p['status']
            self.project_list.setItem(row_idx, 3, QTableWidgetItem(status))
            
            updated_str = str(p['updated_at'])
            try:
                from datetime import datetime
                dt = datetime.strptime(updated_str.split('.')[0], '%Y-%m-%d %H:%M:%S')
                updated_str = dt.strftime('%d/%m/%Y %H:%M:%S')
            except Exception:
                pass
            self.project_list.setItem(row_idx, 4, QTableWidgetItem(updated_str))

    def create_new_project(self):
        dialog = NewProjectDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            project_id = db.create_project(dialog.student_name, dialog.thesis_title, dialog.advisor_name, dialog.advisor_email, dialog.pdf_path)
            self.init_workspace(project_id)

    def edit_selected_project(self):
        selected_items = self.project_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Aviso", "Por favor, selecione um projeto para editar.")
            return
        
        row = selected_items[0].row()
        project_id = int(self.project_list.item(row, 0).text())
        project_data = db.get_project(project_id)
        
        dialog = EditProjectDialog(project_data, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            db.update_project_details(project_id, dialog.student_name, dialog.thesis_title, dialog.advisor_name, dialog.advisor_email)
            self.populate_dashboard()

    def open_selected_project(self):
        selected_items = self.project_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Aviso", "Por favor, selecione um projeto para abrir.")
            return
        
        row = selected_items[0].row()
        project_id = int(self.project_list.item(row, 0).text())
        self.init_workspace(project_id)

    def delete_selected_project(self):
        selected_items = self.project_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Aviso", "Por favor, selecione um projeto para excluir.")
            return
            
        row = selected_items[0].row()
        project_id = int(self.project_list.item(row, 0).text())
        
        reply = QMessageBox.question(self, "Excluir Projeto", 
                                     "Tem certeza de que deseja excluir este projeto e todas as suas anotações?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_project(project_id)
            self.populate_dashboard()

    def toggle_selected_project_status(self):
        selected_items = self.project_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Aviso", "Por favor, selecione um projeto para alterar o status.")
            return
            
        row = selected_items[0].row()
        project_id = int(self.project_list.item(row, 0).text())
        current_status = self.project_list.item(row, 3).text()
        
        new_status = "Concluído" if current_status == "Em Andamento" else "In Progress"
        
        db.update_project_status(project_id, new_status)
        self.populate_dashboard()

    def save_draft_and_exit(self):
        if self.current_project_id:
            self.save_general_opinion()
            db.touch_project(self.current_project_id)
        self.init_dashboard()

    def init_workspace(self, project_id):
        self.current_project_id = project_id
        self.editing_annotation_id = None
        project_data = db.get_project(project_id)
        
        if not project_data:
            QMessageBox.critical(self, "Erro", "Projeto não encontrado no banco de dados.")
            self.init_dashboard()
            return

        self.clear_layout(self.main_layout)
        
        workspace_widget = QWidget()
        layout = QVBoxLayout(workspace_widget)
        
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # ==========================================
        # LEFT COLUMN: PDF Viewer 
        # ==========================================
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        
        pdf_toolbar = QHBoxLayout()
        btn_zoom_in = QPushButton("🔍+")
        btn_zoom_in.setObjectName("btn_icon_blue")
        btn_zoom_in.setToolTip("Ampliar")
        
        btn_zoom_out = QPushButton("🔍-")
        btn_zoom_out.setObjectName("btn_icon_blue")
        btn_zoom_out.setToolTip("Reduzir")
        
        btn_prev = QPushButton("⬅️")
        btn_prev.setObjectName("btn_icon_orange")
        btn_prev.setToolTip("Página Anterior")
        
        btn_next = QPushButton("➡️")
        btn_next.setObjectName("btn_icon_orange")
        btn_next.setToolTip("Próxima Página")
        
        self.page_nav_input = QLineEdit()
        self.page_nav_input.setFixedWidth(50)
        self.page_nav_input.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_total_pages = QLabel(" / X")
        
        btn_go_page = QPushButton("Ir")
        btn_go_page.setObjectName("btn_icon_blue")
        btn_go_page.setToolTip("Ir para página")
        
        pdf_toolbar.addWidget(btn_zoom_in)
        pdf_toolbar.addWidget(btn_zoom_out)
        pdf_toolbar.addWidget(btn_prev)
        pdf_toolbar.addWidget(btn_next)
        pdf_toolbar.addWidget(QLabel("Página: "))
        pdf_toolbar.addWidget(self.page_nav_input)
        pdf_toolbar.addWidget(self.lbl_total_pages)
        pdf_toolbar.addWidget(btn_go_page)
        pdf_toolbar.addStretch()
        
        left_layout.addLayout(pdf_toolbar)
        
        self.pdf_viewer = PDFViewer()
        self.pdf_viewer.textSelected.connect(self.handle_text_captured)
        left_layout.addWidget(self.pdf_viewer)
        
        btn_zoom_in.clicked.connect(self.pdf_viewer.zoom_in)
        btn_zoom_out.clicked.connect(self.pdf_viewer.zoom_out)
        btn_prev.clicked.connect(self.handle_prev_page)
        btn_next.clicked.connect(self.handle_next_page)
        self.page_nav_input.editingFinished.connect(self.go_to_page)
        btn_go_page.clicked.connect(self.go_to_page)
        
        # ==========================================
        # RIGHT COLUMN: Annotation Panel 
        # ==========================================
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        
        info_label = QLabel(f"<b>Aluno:</b> {project_data['student_name']} | <b>Título:</b> {project_data['thesis_title']}")
        info_label.setWordWrap(True)
        right_layout.addWidget(info_label)
        
        form_layout = QFormLayout()
        
        self.category_tabs = QTabWidget()
        self.category_tabs.addTab(QWidget(), "Apontamentos")
        self.category_tabs.addTab(QWidget(), "Dúvida")
        self.category_tabs.addTab(QWidget(), "Sugestão")
        self.category_tabs.currentChanged.connect(self.update_tab_color)
        
        right_layout.addWidget(self.category_tabs)
        
        self.page_input = QLineEdit()
        self.page_input.setPlaceholderText("Preenchido automaticamente da seleção")
        form_layout.addRow("Página:", self.page_input)
        
        self.excerpt_input = QPlainTextEdit()
        self.excerpt_input.setPlaceholderText("Trecho de texto capturado automaticamente...")
        self.excerpt_input.setMaximumHeight(80)
        form_layout.addRow("Trecho:", self.excerpt_input)
        
        self.notes_input = QPlainTextEdit()
        self.notes_input.setPlaceholderText("Digite sua observação ou sugestão aqui...")
        self.notes_input.setMaximumHeight(100)
        form_layout.addRow("Observação:", self.notes_input)
        
        self.btn_save_ann = QPushButton("Salvar Anotação")
        self.btn_save_ann.setObjectName("btn_orange")
        self.btn_save_ann.clicked.connect(self.save_annotation)
        
        self.btn_cancel_edit = QPushButton("Cancelar Edição")
        self.btn_cancel_edit.setObjectName("btn_red")
        self.btn_cancel_edit.clicked.connect(self.cancel_edit_annotation)
        self.btn_cancel_edit.setVisible(False)
        
        btn_ann_layout = QHBoxLayout()
        btn_ann_layout.addWidget(self.btn_save_ann)
        btn_ann_layout.addWidget(self.btn_cancel_edit)
        
        form_layout.addRow("", btn_ann_layout)
        
        right_layout.addLayout(form_layout)
        
        ann_label = QLabel("<b>Anotações Salvas para este Projeto</b>")
        right_layout.addWidget(ann_label)
        
        self.ann_list = QTableWidget(0, 4)
        self.ann_list.setHorizontalHeaderLabels(["ID", "Página", "Categoria", "Trecho"])
        self.ann_list.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.ann_list.hideColumn(0)
        self.ann_list.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.ann_list.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        right_layout.addWidget(self.ann_list)
        
        ann_actions_layout = QHBoxLayout()
        btn_edit_ann = QPushButton("Editar Selecionada")
        btn_edit_ann.setObjectName("btn_blue")
        btn_edit_ann.clicked.connect(self.load_annotation_for_editing)
        
        btn_delete_ann = QPushButton("Excluir Selecionada")
        btn_delete_ann.setObjectName("btn_red")
        btn_delete_ann.clicked.connect(self.delete_selected_annotation)
        
        ann_actions_layout.addWidget(btn_edit_ann)
        ann_actions_layout.addWidget(btn_delete_ann)
        right_layout.addLayout(ann_actions_layout)
        
        lbl_general = QLabel("<b>Parecer Geral de Membro da Banca Avaliadora</b>")
        right_layout.addWidget(lbl_general)
        self.general_opinion = QPlainTextEdit()
        self.general_opinion.setMaximumHeight(100)
        self.general_opinion.setPlaceholderText("Escreva seu parecer geral sobre o TCC aqui...")
        
        if project_data['general_opinion']:
            self.general_opinion.setPlainText(project_data['general_opinion'])
        
        self.general_opinion.textChanged.connect(self.autosave_timer.start)
        right_layout.addWidget(self.general_opinion)
        
        footer_btn_layout = QHBoxLayout()
        btn_export = QPushButton("Exportar Relatório")
        btn_export.setObjectName("btn_blue")
        btn_export.clicked.connect(self.open_export_dialog)
        
        btn_email = QPushButton("Enviar por E-mail")
        btn_email.setObjectName("btn_orange")
        btn_email.clicked.connect(self.send_reports_by_email)
        
        btn_save_exit = QPushButton("Salvar Rascunho e Sair")
        btn_save_exit.setObjectName("btn_orange")
        btn_save_exit.clicked.connect(self.save_draft_and_exit)
        
        btn_validate = QPushButton("Validar Citações")
        btn_validate.setObjectName("btn_green")
        btn_validate.clicked.connect(self.run_citation_validation)
        
        footer_btn_layout.addWidget(btn_validate)
        footer_btn_layout.addWidget(btn_export)
        footer_btn_layout.addWidget(btn_email)
        footer_btn_layout.addWidget(btn_save_exit)
        right_layout.addLayout(footer_btn_layout)
        
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        
        # Define stretch factors and initial sizes to force 65% / 35% ratio
        splitter.setStretchFactor(0, 65)
        splitter.setStretchFactor(1, 35)
        splitter.setSizes([6500, 3500])
        
        layout.addWidget(splitter)
        self.main_layout.addWidget(workspace_widget)
        
        self.update_tab_color(0)
        
        if project_data['pdf_path']:
            success = self.pdf_viewer.load_pdf(project_data['pdf_path'])
            if success:
                self.update_page_label()
            else:
                QMessageBox.warning(self, "Erro", f"Falha ao carregar PDF de:\n{project_data['pdf_path']}")
        
        self.populate_annotations()

    def handle_prev_page(self):
        self.pdf_viewer.prev_page()
        self.update_page_label()

    def handle_next_page(self):
        self.pdf_viewer.next_page()
        self.update_page_label()

    def update_page_label(self):
        if hasattr(self, 'pdf_viewer') and self.pdf_viewer.pdf_document:
            current = self.pdf_viewer.current_page_idx + 1
            total = len(self.pdf_viewer.pdf_document)
            if hasattr(self, 'page_nav_input'):
                self.page_nav_input.setText(str(current))
            if hasattr(self, 'lbl_total_pages'):
                self.lbl_total_pages.setText(f" / {total}")

    def go_to_page(self):
        if not hasattr(self, 'pdf_viewer') or not self.pdf_viewer.pdf_document:
            return
        
        try:
            page_num = int(self.page_nav_input.text())
            total = len(self.pdf_viewer.pdf_document)
            if 1 <= page_num <= total:
                self.pdf_viewer.current_page_idx = page_num - 1
                self.pdf_viewer.render_page()
                self.update_page_label()
            else:
                self.update_page_label()
        except ValueError:
            self.update_page_label()

    def handle_text_captured(self, text, page_num):
        self.page_input.setText(str(page_num))
        self.excerpt_input.setPlainText(text)
        self.notes_input.setFocus()

    def save_annotation(self):
        if not self.current_project_id: return
        
        page_str = self.page_input.text()
        excerpt = self.excerpt_input.toPlainText().strip()
        notes = self.notes_input.toPlainText().strip()
        category = self.category_tabs.tabText(self.category_tabs.currentIndex())
        
        if not page_str.isdigit():
            QMessageBox.warning(self, "Validação", "A página deve ser um número inteiro.")
            return
            
        page = int(page_str)
        
        if self.editing_annotation_id:
            db.update_annotation(self.editing_annotation_id, page, excerpt, category, notes)
            self.editing_annotation_id = None
            self.btn_save_ann.setText("Salvar Anotação")
            self.btn_cancel_edit.setVisible(False)
        else:
            db.add_annotation(self.current_project_id, page, excerpt, category, notes)
        
        self.page_input.clear()
        self.excerpt_input.clear()
        self.notes_input.clear()
        
        self.populate_annotations()
        
    def cancel_edit_annotation(self):
        self.editing_annotation_id = None
        self.btn_save_ann.setText("Salvar Anotação")
        self.btn_cancel_edit.setVisible(False)
        self.page_input.clear()
        self.excerpt_input.clear()
        self.notes_input.clear()
        
    def load_annotation_for_editing(self):
        selected_items = self.ann_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Aviso", "Por favor, selecione uma anotação para editar.")
            return
            
        row = selected_items[0].row()
        annotation_id = int(self.ann_list.item(row, 0).text())
        
        anns = db.get_annotations(self.current_project_id)
        ann = next((a for a in anns if a['id'] == annotation_id), None)
        
        if ann:
            self.editing_annotation_id = annotation_id
            self.page_input.setText(str(ann['page_number']))
            self.excerpt_input.setPlainText(ann['selected_text'])
            self.notes_input.setPlainText(ann['professor_notes'])
            
            # Find tab index
            for i in range(self.category_tabs.count()):
                if self.category_tabs.tabText(i) == ann['category']:
                    self.category_tabs.setCurrentIndex(i)
                    break
                    
            self.btn_save_ann.setText("Atualizar Anotação")
            self.btn_cancel_edit.setVisible(True)

    def delete_selected_annotation(self):
        selected_items = self.ann_list.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "Aviso", "Por favor, selecione uma anotação para excluir.")
            return
            
        row = selected_items[0].row()
        annotation_id = int(self.ann_list.item(row, 0).text())
        
        reply = QMessageBox.question(self, "Excluir Anotação", 
                                     "Tem certeza de que deseja excluir esta anotação?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            db.delete_annotation(annotation_id)
            if self.editing_annotation_id == annotation_id:
                self.cancel_edit_annotation()
            self.populate_annotations()

    def populate_annotations(self):
        if not self.current_project_id: return
        self.ann_list.setRowCount(0)
        anns = db.get_annotations(self.current_project_id)
        for row_idx, ann in enumerate(anns):
            self.ann_list.insertRow(row_idx)
            self.ann_list.setItem(row_idx, 0, QTableWidgetItem(str(ann['id'])))
            self.ann_list.setItem(row_idx, 1, QTableWidgetItem(str(ann['page_number'])))
            self.ann_list.setItem(row_idx, 2, QTableWidgetItem(ann['category']))
            
            excerpt = ann['selected_text']
            if len(excerpt) > 40:
                excerpt = excerpt[:37] + "..."
            self.ann_list.setItem(row_idx, 3, QTableWidgetItem(excerpt))

    def save_general_opinion(self):
        if self.current_project_id:
            text = self.general_opinion.toPlainText()
            db.update_general_opinion(self.current_project_id, text)

    def run_citation_validation(self):
        if not self.current_project_id: return
        project_data = db.get_project(self.current_project_id)
        pdf_path = project_data['pdf_path']
        
        if not pdf_path or not os.path.exists(pdf_path):
            QMessageBox.warning(self, "Erro", "PDF não encontrado.")
            return
            
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        try:
            results = validador_citacoes.run_validation(pdf_path)
            QApplication.restoreOverrideCursor()
            
            if not results:
                QMessageBox.information(self, "Aviso", "Nenhuma citação processada ou falha na extração.")
                return
                
            dialog = CitationReportDialog(results, self)
            dialog.exec()
        except Exception as e:
            QApplication.restoreOverrideCursor()
            QMessageBox.critical(self, "Erro", f"Ocorreu um erro durante a validação:\n{e}")

    def update_tab_color(self, index):
        colors = ["rgba(239, 68, 68, 0.2)", "rgba(234, 179, 8, 0.2)", "rgba(59, 130, 246, 0.2)"]
        borders = ["rgba(239, 68, 68, 0.6)", "rgba(234, 179, 8, 0.6)", "rgba(59, 130, 246, 0.6)"]
        self.category_tabs.setStyleSheet(
            f"QTabBar::tab:selected {{ background: {colors[index]}; border-bottom: 2px solid {borders[index]}; font-weight: bold; color: white; }}"
        )

    def open_export_dialog(self):
        if not self.current_project_id: return
        
        dialog = ExportDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            evaluator = dialog.evaluator_name
            date_str = dialog.date_str
            fmt = dialog.export_format
            
            project_data = db.get_project(self.current_project_id)
            default_name = f"Revisao_{project_data['student_name'].replace(' ', '_')}.{fmt}"
            
            filter_str = "Documentos PDF (*.pdf)" if fmt == "pdf" else "Documentos do Word (*.docx)"
            file_path, _ = QFileDialog.getSaveFileName(self, "Salvar Relatório", default_name, filter_str)
            
            if not file_path: return
            
            self.save_general_opinion()
            
            if fmt == "pdf":
                self.generate_pdf(file_path, project_data, evaluator, date_str)
            else:
                self.generate_docx(file_path, project_data, evaluator, date_str)

    def generate_pdf(self, file_path, project_data, evaluator, date_str, show_msg=True):
        anns = db.get_annotations(self.current_project_id)
        
        html = f"""
        <html>
        <head>
            <style>
                body, div, p {{ font-family: Arial, sans-serif; font-size: 12pt; line-height: 1.6; color: #000; }}
                h1 {{ color: #000; text-align: center; border-bottom: 2px solid #000; padding-bottom: 10pt; font-size: 20pt; }}
                h2 {{ color: #000; margin-top: 30pt; font-size: 14pt; }}
                .info {{ font-size: 14pt; margin-bottom: 20pt; }}
                .notes {{ margin-bottom: 20pt; }}
                .signature-section {{ margin-top: 50pt; text-align: center; }}
                .cursive-signature {{ font-family: 'Lucida Handwriting', 'Brush Script MT', cursive; font-size: 12pt; color: #000; margin-bottom: 5pt; }}
                .signature-line {{ border-top: 1px solid #000; width: 300pt; margin: 0 auto; margin-top: 10pt; padding-top: 5pt; }}
            </style>
        </head>
        <body>
            <h1>Relatório de Revisão de TCC</h1>
            <div class='info'>
                <b>Aluno(a):</b> {project_data['student_name']}<br>
                <b>Título:</b> {project_data['thesis_title']}<br>
                <b>Avaliador(a):</b> {evaluator if evaluator else 'Professor(a) / Banca Examinadora'}
            </div>
        """
        
        for ann in anns:
            html += f"<h2>Página {ann['page_number']} - Categoria: {ann['category']}</h2>"
            html += f"<div><b>Trecho Citado:</b> \"{ann['selected_text']}\"</div>"
            html += f"<div class='notes'><b>Observações e Sugestões:</b><br>{ann['professor_notes'].replace(chr(10), '<br>')}</div>"
        
        html += "<h2>Parecer Geral</h2>"
        html += f"<div>{(project_data['general_opinion'] or 'Sem parecer geral.').replace(chr(10), '<br>')}</div>"
        
        html += f"""
            <div class='signature-section'>
                <div class='cursive-signature'>{evaluator if evaluator else ''}</div>
                <div class='signature-line'>{evaluator if evaluator else 'Assinatura'}</div>
                <div>{date_str}</div>
            </div>
        </body>
        </html>
        """
        
        document = QTextDocument()
        document.setHtml(html)
        
        writer = QPdfWriter(file_path)
        writer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
        writer.setPageMargins(QMarginsF(15, 15, 15, 15), QPageLayout.Unit.Millimeter)
        
        document.print(writer)
        if show_msg:
            QMessageBox.information(self, "Sucesso", f"Relatório PDF exportado com sucesso para:\n{file_path}")

    def generate_docx(self, file_path, project_data, evaluator, date_str, show_msg=True):
        try:
            import docx
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = docx.Document()
            
            doc.add_heading('Relatório de Revisão de Trabalho de Conclusão de Curso (TCC)', 0)
            
            p_info = doc.add_paragraph()
            r_student = p_info.add_run(f"Aluno(a): {project_data['student_name']}\n")
            r_student.font.size = Pt(14)
            r_student.bold = True
            
            r_title = p_info.add_run(f"Título: {project_data['thesis_title']}\n")
            r_title.font.size = Pt(14)
            r_title.bold = True
            
            r_evaluator = p_info.add_run(f"Avaliador(a): {evaluator if evaluator else 'Professor(a) / Banca Examinadora'}")
            r_evaluator.font.size = Pt(14)
            r_evaluator.bold = True
            
            doc.add_paragraph() 
            
            anns = db.get_annotations(self.current_project_id)
            for ann in anns:
                doc.add_heading(f"Página {ann['page_number']} - Categoria: {ann['category']}", level=2)
                
                doc.add_paragraph(f"Trecho Citado: \"{ann['selected_text']}\"")
                doc.add_paragraph(f"Observações e Sugestões:\n{ann['professor_notes']}")
                doc.add_paragraph() 
            
            doc.add_heading('Parecer Geral de Membro da Banca Avaliadora', level=1)
            doc.add_paragraph(project_data['general_opinion'] or "Sem parecer geral.")
            
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()
            
            p_sig = doc.add_paragraph()
            p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if evaluator:
                r_sig = p_sig.add_run(evaluator)
                r_sig.font.name = 'Lucida Handwriting'
                r_sig.font.size = Pt(12)
            
            p_line = doc.add_paragraph()
            p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_line = p_line.add_run("_________________________________________________")
            
            p_name = doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_name.add_run(evaluator if evaluator else 'Assinatura')
            
            p_date = doc.add_paragraph()
            p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_date.add_run(date_str)
            
            doc.save(file_path)
            if show_msg:
                QMessageBox.information(self, "Sucesso", f"Relatório DOCX exportado com sucesso para:\n{file_path}")
            
        except ImportError:
            QMessageBox.critical(self, "Erro", "python-docx não está instalado. Por favor, rode: pip install python-docx")
        except Exception as e:
            QMessageBox.critical(self, "Erro", f"Falha ao exportar relatório:\n{str(e)}")

    def send_reports_by_email(self):
        if not self.current_project_id: return
        project_data = db.get_project(self.current_project_id)
        
        advisor_email = project_data.get('advisor_email', '').strip()
        if not advisor_email:
            QMessageBox.warning(self, "Aviso", "O e-mail do orientador não está configurado neste projeto. Edite o projeto para adicionar o e-mail.")
            return
            
        dialog = ExportDialog(self)
        dialog.setWindowTitle("Configurar Envio de E-mail")
        if dialog.exec() == QDialog.DialogCode.Accepted:
            evaluator = dialog.evaluator_name
            date_str = dialog.date_str
            fmt = dialog.export_format
            
            QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
            
            try:
                import tempfile
                temp_dir = tempfile.gettempdir()
                general_report_path = os.path.join(temp_dir, f"Parecer_{project_data['student_name'].replace(' ', '_')}.{fmt}")
                if fmt == "pdf":
                    self.generate_pdf(general_report_path, project_data, evaluator, date_str, show_msg=False)
                else:
                    self.generate_docx(general_report_path, project_data, evaluator, date_str, show_msg=False)
                    
                results = validador_citacoes.run_validation(project_data['pdf_path'])
                citation_report_path = os.path.join(temp_dir, f"Citacoes_{project_data['student_name'].replace(' ', '_')}.xlsx")
                import pandas as pd
                df = pd.DataFrame(results)
                df.to_excel(citation_report_path, index=False)
                
                self._send_email_smtp(advisor_email, general_report_path, citation_report_path, project_data)
                
                QApplication.restoreOverrideCursor()
                QMessageBox.information(self, "Sucesso", "E-mail enviado com sucesso ao orientador!")
                
            except Exception as e:
                QApplication.restoreOverrideCursor()
                QMessageBox.critical(self, "Erro", f"Falha ao enviar e-mail:\n{str(e)}")
                
    def _send_email_smtp(self, recipient_email, report1_path, report2_path, project_data):
        smtp_server = os.environ.get("SMTP_SERVER", "")
        smtp_port = int(os.environ.get("SMTP_PORT", "587"))
        smtp_user = os.environ.get("SMTP_USER", "")
        smtp_password = os.environ.get("SMTP_PASSWORD", "")
        
        if not smtp_server or not smtp_user or not smtp_password:
            raise Exception("Credenciais SMTP não configuradas no arquivo .env.")
            
        msg = EmailMessage()
        msg['Subject'] = f"Relatórios de Revisão de TCC - {project_data['student_name']}"
        msg['From'] = smtp_user
        msg['To'] = recipient_email
        msg.set_content(f"Prezado(a) Orientador(a),\n\nSegue em anexo o Parecer Geral e o Relatório de Validação de Citações referente ao TCC de {project_data['student_name']}.\n\nTítulo: {project_data['thesis_title']}\n\nAtenciosamente,\nAssistente de Revisão de TCC.")
        
        import mimetypes
        for file_path in [report1_path, report2_path]:
            if not os.path.exists(file_path): continue
            ctype, encoding = mimetypes.guess_type(file_path)
            if ctype is None or encoding is not None:
                ctype = 'application/octet-stream'
            maintype, subtype = ctype.split('/', 1)
            with open(file_path, 'rb') as fp:
                msg.add_attachment(fp.read(), maintype=maintype, subtype=subtype, filename=os.path.basename(file_path))
                
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.send_message(msg)

    def clear_layout(self, layout):
        while layout.count():
            child = layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
            elif child.layout():
                self.clear_layout(child.layout())

if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setStyleSheet(DARK_GLASS_QSS)
    window = TCCAssistantApp()
    window.show()
    sys.exit(app.exec())
